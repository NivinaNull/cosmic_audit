# -*-coding:utf-8 -*-

"""
# File       : vector_based_sim.py
# Time       ：2020/9/9 17:01
# Author     ：nivinanull@163.com
# version    ：python 3.6
"""

import re
import sys
import copy
import numpy as np
import pandas as pd
import datetime
import jieba
import jieba.posseg as pseg
from collections import OrderedDict
from docx import Document
from docx.shared import Pt,RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


pd.set_option('display.max_columns', 20)   #给最大列设置为20列
pd.set_option('display.max_rows', 10)   #设置最大可见10行
jieba.load_userdict("./vector_data/custom_dict.txt")


def get_corpus_data(cosmic_info=pd.DataFrame(), noncosmic_info=pd.DataFrame(), process_type=''):
    # stopwords = [line.decode('utf-8').strip() for line in codecs.open('./data/stopwords.txt', 'rb').readlines()]
    stop_property = ['nr', 'o', 'p', 'q', 'r', 'x', 'y', 'w', 'uj', 'm', 'un', 'c']

    stop_words_check = dict.fromkeys(stop_property)

    def num_process(w):
        if w.flag in stop_property:
            if not stop_words_check[w.flag]:
                stop_words_check[w.flag] = [w.flag]
            else:
                stop_words_check[w.flag].append(w.word)
            return False
        else:
            return True

    def segment(row):
        if row:
            year_list = [str(i + datetime.datetime.now().year) for i in range(-20, 21)]
            year_related = ['\d*' + y + '\d*' for y in year_list]
            year_related_str = '|'.join(year_related)
            row = re.sub(year_related_str + '|\d\.|\d、|\d,|\d，|\W+', ' ', row).replace('_', ' ')
            # row = re.sub('\d\.|\d、|\d|\n|\t', '', row)
            jieba.load_userdict('./data/my_dict.txt')
            # words = jieba.lcut(row)
            # words = [i for i in words if i.strip() and num_process(i)]
            words = pseg.lcut(row)
            words = [i.word for i in words if num_process(i)]
            # words = [i.word for i in words if lambda i: True if i.flag not in stop_property else False]

        else:
            words = []
        # return ' '.join(words)
        return words

    def to_str(attri_row):
        joint_info = ''
        for c in range(len(attri_row.work_detail)):
            if not (attri_row.work_detail[c].strip() == '无' or attri_row.work_detail[c].strip() == ''):
                joint_info = attri_row.work_detail[c]
                if not (attri_row.work_name[c].strip() == '无' or attri_row.work_name[c].strip() == ''):
                    joint_info = joint_info + ' ' + attri_row.work_name[c]
                if not (attri_row.work_cat[c].strip() == '无' or attri_row.work_cat[c].strip() == ''):
                    joint_info = joint_info + ' ' + attri_row.work_cat[c] + ' '
        return joint_info

    cosmic_inner = copy.deepcopy(cosmic_info)
    noncosmic_inner = copy.deepcopy(noncosmic_info)
    if not noncosmic_inner.empty:
        noncosmic_inner['joint_info'] = noncosmic_inner[['work_cat', 'work_name', 'work_detail']].apply(to_str, axis=1)
        noncosmic_inner['joint_info'] = noncosmic_inner['joint_info'] + noncosmic_inner['requirement_name']
        noncosmic_inner['joint_info'] = noncosmic_inner['joint_info'].map(segment)
    else:
        return False, 'ERROR:noncosmic信息为空'

    if not cosmic_inner.empty:
        cosmic_inner['joint_info'] = ''
        for i in ['project_name', 'requirement_name', 'requirement_detail', 'coding_requirement']:
            cosmic_inner['joint_info'] = cosmic_inner['joint_info'] + ' ' + cosmic_inner[i]
        cosmic_inner['joint_info'] = cosmic_inner['joint_info'].map(segment)
    else:
        return False, 'ERROR:cosmic信息为空'

    cosmic_inner.drop(columns=['project_name', 'requirement_name', 'requirement_detail', 'advocator', 'days_spent', 'coding_requirement'], inplace=True)
    noncosmic_inner.drop(columns=['project_name', 'requirement_name', 'work_cat', 'work_detail', 'work_name', 'days_spent'], inplace=True)
    return True, [cosmic_inner, noncosmic_inner]


def load_vector():
    # wv_from_text = KeyedVectors.load_word2vec_format('./vector_data/70000-small.txt', binary=False)
    # print('wv_from_text.index2entity ' * 10 + '\n', len(wv_from_text.index2entity))
    # print('wv_from_text.index2word ' * 10 + '\n', len(wv_from_text.index2word))

    embedding_index = {}
    f = open('./vector_data/70000-small.txt', encoding='utf8')
    for index, line in enumerate(f):
        if index == 0:
            continue
        values = line.split(' ')
        word = values[0]
        coefs = np.asarray(values[1:], dtype='float32')
        embedding_index[word] = coefs
    f.close()
    return embedding_index


def r2vec(cosmic_corpus, noncosmic_corpus, vectors):
    # cosmic_corpus
    cosmic_corpus['text_vec'] = [np.zeros(200, dtype=np.float32) for i in range(cosmic_corpus.shape[0])]
    noncosmic_corpus['text_vec'] = [np.zeros(200, dtype=np.float32) for i in range(noncosmic_corpus.shape[0])]

    def corpus_convert(text):
        text_len = len(text)
        to_caculate = np.zeros(200, dtype=np.float32)
        for w in text:
            if w in vectors:
                to_caculate += vectors[w]
            else:
                text_len -= 1
        # print(to_caculate)
        to_caculate = to_caculate / text_len
        # print(text_len, '\n', to_caculate)
        return to_caculate

    cosmic_corpus['text_vec'] = cosmic_corpus['joint_info'].apply(corpus_convert)
    noncosmic_corpus['text_vec'] = noncosmic_corpus['joint_info'].apply(corpus_convert)

    return cosmic_corpus.drop(columns=['joint_info']), noncosmic_corpus.drop(columns=['joint_info'])


# 模拟新数据
def get_new_req(cosmic, noncosmic):
    # return cosmic.sample(n=10), noncosmic.sample(n=10)
    return cosmic[~(cosmic.projectNo == '321')], noncosmic


def get_sim(his, new, distance='cosine'):
    his.reset_index(drop=True, inplace=True)
    new.reset_index(drop=True, inplace=True)
    np_his = np.zeros((his.shape[0], 200))
    np_new = np.zeros((new.shape[0], 200))
    for i in range(his.shape[0]):
        np_his[i] = his['text_vec'].values[i]
    for i in range(new.shape[0]):
        np_new[i] = new['text_vec'].values[i]

    # inner_product =  rn * rh = (rn * 200) * (rh * 200)T
    inner_product = np.dot(np_new, np_his.transpose())
    # norm_new = (rn, )
    norm_new = np.sqrt(np.multiply(np_new, np_new).sum(axis=1))
    # norm_new = rn * 1
    norm_new = norm_new[:, np.newaxis]
    # norm_his = (rh, )
    norm_his = np.sqrt(np.multiply(np_his, np_his).sum(axis=1))
    # norm_his = rh * 1
    norm_his = norm_his[:, np.newaxis]
    # np.dot(norm_new, norm_his.transpose()) = rn * rh = (rn * 1) * (rh * 1)T
    cosine_distance = np.divide(inner_product, np.dot(norm_new, norm_his.transpose()))
    # print('cosine_distance.shape' * 20 + '\n', cosine_distance.shape)
    pos = np.where(cosine_distance >= 0.98)
    sim_value = cosine_distance[pos]
    # print('pos_ ' * 20 + '\n', pos)
    # print('len(sim_value) ' * 20 + '\n', len(sim_value))
    sim_index = {}
    for j in range(len(sim_value)):
        new_key = (new.iloc[pos[0][j]].batch, new.iloc[pos[0][j]].projectNo, new.iloc[pos[0][j]].requirementNO)
        his_key = (his.iloc[pos[1][j]].batch, his.iloc[pos[1][j]].projectNo, his.iloc[pos[1][j]].requirementNO)
        if his_key == new_key:
            continue
        if new_key not in sim_index:
            sim_index[new_key] = {}
            sim_index[new_key][his_key] = sim_value[j]
        else:
            # print('check_it_out ' * 20 + '\n', sim_index[new_key])
            sim_index[new_key][his_key] = sim_value[j]
            # print('check_it_out ' * 20 + '\n', sim_index[new_key])
    return sim_index


def show(result_dict, raw_info, new, flag):
    print(result_dict)
    if result_dict:
        for k in result_dict.keys():
            req_detail = new[(new.batch == k[0]) & (new.projectNo == k[1]) & (new.requirementNO == k[2])]
            # print('check_it_out ' * 20 + '\n', req_detail)
            print('发现' + flag + '需求批次' + k[0] + '项目编号' + k[1] + '需求编号' + str(k[2]) + '\n', '需求详细信息：',
                  req_detail.to_json(force_ascii=False, orient='split', index=False), '相似需求如下：')
            # print('result_dict[k]111' * 5, result_dict[k])
            result_dict[k] = sorted(result_dict[k].items(), key=lambda x: x[1], reverse=True)
            # print('result_dict[k]222' * 5, result_dict[k])
            for ki in result_dict[k]:
                sim_detail = raw_info[(raw_info.batch == ki[0][0]) & (raw_info.projectNo == ki[0][1]) & (
                            raw_info.requirementNO == ki[0][2])]
                print('相似需求批次' + ki[0][0] + '项目编号' + ki[0][1] + '需求编号' + str(ki[0][2]) + '\n',
                      sim_detail.to_json(force_ascii=False, orient='split', index=False))
                print('相似度为：', ki[1])
    else:
        print('未发现相似' + flag + '需求')

def result_persist(result_dict, raw_info, new, related_no, flag):
    # print('right in result_persist  ' * 20)

    def df2dict(req_df):
        req_ser = req_df.iloc[0]
        req_dict = req_ser.to_dict()
        des = OrderedDict([
            ('batch', '需求批次'),
            ('projectNo', '项目编号'),
            ('project_name', '项目名称'),
            ('requirementNO', '需求编号'),
            ('requirement_name', '需求名称'),
            ('days_spent', '实际消耗时间'),
        ])
        if flag == 'cosmic':
            des['advocator'] = '需求提出人'
            des['requirement_detail'] = '需求详细信息'
            des['coding_requirement'] = '代码开发信息'
        else:
            des['work_cat'] = '工作类别'
            des['work_name'] = '工作名称'
            des['work_detail'] = '工作详情'
        for n in des:
            c = des[n] + '：' + str(req_dict[n])
            # print('here ' * 50 + '\n',c , '\n length:', len(c))
            if len(c) <= 100:
                des[n] = c
            else:
                des[n] = c[:35] + '\t ...'
                # print('len(c) > 100 ' * 20, des[n])
        # return '\n'.join(des.values())
        return des

    def in_paragraph(req_dict, DOCobj, sim_value=''):
        if not sim_value:
            mes = '#' * 50 + '\n' + '原需求如下：\n'
        else:
            mes = '相似需求如下：\n'
        p = DOCobj.add_paragraph()
        p.add_run(mes).bold = True
        mes = '\n'.join([req_dict[key] for key in ['batch', 'projectNo', 'project_name']]) + '\n'
        p.add_run(mes)
        mes = '\n'.join([req_dict[key] for key in ['requirementNO', 'requirement_name']]) + '\n'
        pr2 = p.add_run(mes)
        pr2.font.color.rgb = RGBColor(0xff, 0, 0)
        [req_dict.pop(key) for key in ['batch', 'projectNo', 'project_name', 'requirementNO', 'requirement_name']]
        mes = '\n'.join(req_dict.values()) + '\n'
        p.add_run(mes)
        if sim_value:
            p.add_run('相似度为：' + sim_value).font.color.rgb = RGBColor(255, 127, 0)

    if result_dict:
        doc = Document()
        doc.styles['Normal'].font.name = '宋体'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        doc.add_heading('TF-IDF Similarity Analyse', level=0)
        doc.add_heading(flag + '的相似需求情况如下：')
        pt = doc.add_paragraph()
        pt.add_run('共计项目个数：' + str(project_no) + '\n'
                   '共计需求个数：' + str(req_no) + '\n'
                   '发现有相似关系的需求：' + str(related_no) + '\n')
        pt.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for k in result_dict.keys():
            # mes += '发现需求批次：{}\t项目编号：{}\t需求编号：{}'.format(k[0], k[1], str(k[2]))
            req_detail = new[(new.batch == k[0]) & (new.projectNo == k[1]) & (new.requirementNO == k[2])]
            req_detail_dict = df2dict(req_detail)
            in_paragraph(req_detail_dict, doc)

            # mes = mes + '#' * 100 + '\n' + json.dumps(req_detail_des, ensure_ascii=False)
            result_dict[k] = sorted(result_dict[k].items(), key=lambda x: x[1], reverse=True)
            sim_req_detail = []
            print(result_dict[k])
            for t in result_dict[k]:
                sim_req = raw_info[(raw_info.batch == t[0][0]) & (raw_info.projectNo == t[0][1]) & (raw_info.requirementNO == t[0][2])]
                sim_req = df2dict(sim_req)
                in_paragraph(sim_req, doc, str(t[1]))

            # mes = mes + '\n'.join(sim_req_detail) + '\n'
        # print("in_result_persist_3" * 20 + flag + "\n", mes)
        # doc.add_paragraph(mes)
        doc.save('./Doc/' + flag + '_TF-IDF.docx')
    else:
        print('未发现相似需求')


# load_vector()
if __name__ == '__main__':
    raw_cosmic = pd.read_pickle('./data/cosmic_info.pkl').fillna('')
    raw_noncosmic = pd.read_pickle('./data/noncosmic_info.pkl').fillna('')
    IF_SUCCESS, info = get_corpus_data(raw_cosmic, raw_noncosmic)
    if not IF_SUCCESS:
        print(info)
        sys.exit()
    w2v = load_vector()
    cosmic_vec, noncosmic_vec = r2vec(info[0], info[1].head(1), w2v)
    new_cosmic, new_noncosmic = get_new_req(raw_cosmic, raw_noncosmic)
    IF_SUCCESS, new = get_corpus_data(new_cosmic, new_noncosmic)
    if not IF_SUCCESS:
        print(info)
        sys.exit()
    new_cosmic_vec, new_noncosmic_vec = r2vec(new[0], new[1], w2v)
    # print(new_cosmic_vec.columns, new_noncosmic_vec.columns)
    result = get_sim(cosmic_vec, new_cosmic_vec)
    result_persist(result, raw_cosmic, new_cosmic, 'cosmic')