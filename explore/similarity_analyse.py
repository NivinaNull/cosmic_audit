#-*- coding: utf-8 -*-
# @Time    : 2020/8/4 10:28
# @Author  : Nivina
# @Email   : nivinanull@163.com
# @File    : similarity_analyse.py

import numpy as np
import pandas as pd
import datetime
import copy
import sys
import jieba
from collections import OrderedDict
import jieba.posseg as pseg
import codecs
import re
import json
from gensim import corpora, models
from docx import Document
from docx.shared import Pt,RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from gensim.similarities import SparseMatrixSimilarity


pd.set_option('display.max_columns',20) #给最大列设置为20列
pd.set_option('display.max_rows',10)#设置最大可见10行


def get_new_req(cosmic, noncosmic, flag):
    if flag == 'cosmic':
        return cosmic
        # return cosmic.sample(n=10)
    elif flag == 'noncosmic':
        return noncosmic
        # return noncosmic.sample(n=10)


# result = get_sim(info[FLAG], new_req_corpus, flag)
# 已存需求的分词词料，新需求的分词词料，FLAG
def get_sim(previous_req, new_req, flag):
    # 利用cosmic和非cosmic信息构建的Dictionary
    dictionary = corpora.Dictionary.load_from_text('./data/total.dic')
    # print('new_req.shape：', new_req.shape, 'raw_info.shape:', raw_info.shape, 'previous_req.shape:', previous_req.shape)
    corpus = [dictionary.doc2bow(text) for text in previous_req[flag]]
    tfidf = models.TfidfModel(corpus)
    index = SparseMatrixSimilarity(tfidf[corpus], 4000)

    new = [dictionary.doc2bow(t) for t in new_req['joint_info']]
    sim_dict = {}
    sim = index[new]
    relation_exist = []
    for i in range(new_req.shape[0]):
        key = (new_req['batch'].iloc[i], new_req['projectNo'].iloc[i], new_req['requirementNO'].iloc[i])
        value = {}
        current = {key,}
        # print(key)
        for j in range(len(sim[i])):
            if sim[i][j] >= 0.8:
                inner_key = (previous_req['batch'].iloc[j], previous_req['projectNo'].iloc[j], previous_req['requirementNO'].iloc[j])
                if inner_key == key:
                    continue
                else:
                    value[inner_key] = sim[i][j]
                    current.add(inner_key)
        if value and (current not in relation_exist):
            relation_exist.append(current)
            sim_dict[key] = value

    return sim_dict, len(relation_exist)


def get_corpus_data(cosmic_info=pd.DataFrame(), noncosmic_info=pd.DataFrame(), result_type='total'):
    # stopwords = [line.decode('utf-8').strip() for line in codecs.open('./data/stopwords.txt', 'rb').readlines()]
    stop_property = ['nr', 'o', 'p', 'q', 'r', 'x', 'y', 'w', 'uj', 'm', 'un', 'c']

    stop_words_check = dict.fromkeys(stop_property)

    def num_process(w):
        if w.flag in stop_property:
            if not stop_words_check[w.flag]:
                stop_words_check[w.flag] = [w.flag]
            else:
                stop_words_check[w.flag].append(w.word)
            return False
        else:
            return True

    def segment(row):
        if row:
            year_list = [str(i + datetime.datetime.now().year) for i in range(-20, 21)]
            year_related = ['\d*' + y + '\d*' for y in year_list]
            year_related_str = '|'.join(year_related)
            row = re.sub(year_related_str + '|\d\.|\d、|\d,|\d，|\W+', ' ', row).replace('_', ' ')
            # row = re.sub('\d\.|\d、|\d|\n|\t', '', row)
            jieba.load_userdict('./data/my_dict.txt')
            # words = jieba.lcut(row)
            # words = [i for i in words if i.strip() and num_process(i)]
            words = pseg.lcut(row)
            words = [i.word for i in words if num_process(i) and not re.search('R|r[0-9]\d+', i.word)]
            # words = [i.word for i in words if lambda i: True if i.flag not in stop_property else False]

        else:
            words = []
        # return ' '.join(words)
        return words

    def to_str(attri_row):
        joint_info = ''
        for c in range(len(attri_row.work_detail)):
            if not (attri_row.work_detail[c].strip() == '无' or attri_row.work_detail[c].strip() == ''):
                joint_info = attri_row.work_detail[c]
                if not (attri_row.work_name[c].strip() == '无' or attri_row.work_name[c].strip() == ''):
                    joint_info = joint_info + ' ' + attri_row.work_name[c]
                if not (attri_row.work_cat[c].strip() == '无' or attri_row.work_cat[c].strip() == ''):
                    joint_info = joint_info + ' ' + attri_row.work_cat[c] + ' '
        return joint_info

    cosmic_inner = copy.deepcopy(cosmic_info)
    noncosmic_inner = copy.deepcopy(noncosmic_info)
    if not noncosmic_inner.empty:
        noncosmic_inner['joint_info'] = noncosmic_inner[[ 'work_cat', 'work_name', 'work_detail']].apply(to_str, axis=1 )
        noncosmic_inner['joint_info'] = noncosmic_inner[ 'joint_info'] + noncosmic_inner['requirement_name']
        if result_type == 'noncosmic':
            noncosmic_inner['joint_info'] = noncosmic_inner['joint_info'].map(segment)
            return True, noncosmic_inner
        elif result_type == 'total':
            noncosmic_inner.drop(
                columns=['work_cat', 'work_name', 'work_detail', 'requirement_name', 'project_name', 'days_spent'], inplace=True)
            noncosmic_inner.rename(columns={'joint_info':'noncosmic'}, inplace=True)
        else:
            return False, 'ERROR:您传入了不需要的noncosmic信息，请检查'

    if not cosmic_inner.empty:
        cosmic_inner['joint_info'] = ''
        for i in ['project_name', 'requirement_name', 'requirement_detail', 'coding_requirement']:
            cosmic_inner['joint_info'] = cosmic_inner['joint_info'] + ' ' + str(cosmic_inner[i])
        if result_type == 'cosmic':
            cosmic_inner['joint_info'] = cosmic_inner['joint_info'].map(segment)
            return True, cosmic_inner
        elif result_type == 'total':
            cosmic_inner.rename(columns={'joint_info':'cosmic'}, inplace=True)
            cosmic_inner = cosmic_inner[['batch', 'projectNo', 'requirementNO', 'cosmic']]
            total = pd.merge(cosmic_inner, noncosmic_inner, how='outer', on=['batch', 'projectNo', 'requirementNO']).fillna('')
            total['combo'] = total['cosmic'] + ' ' + total['noncosmic']
            total[['cosmic', 'noncosmic', 'combo']] = total[['cosmic', 'noncosmic', 'combo']].applymap(segment)
            ################################################check the stopwords ##########################################
            for k in stop_words_check.keys():
                with open('./data/stopwords_folder/' + k + '.txt', 'wb') as f:
                    f.write(str(stop_words_check[k]).encode('utf-8'))
            ##############################################################################################################
            return True, total
        else:
            return False, 'ERROR:您传入了不需要的cosmic信息，请检查'
    return False, 'ERROR：未发现有效数据'


def show(result_dict, raw_info, new, flag):
    # print(result_dict)
    if result_dict:
        for k in result_dict.keys():
            req_detail = new[(new.batch == k[0]) & (new.projectNo == k[1]) & (new.requirementNO == k[2])]
            print('发现' + flag + '需求批次' + k[0] + '项目编号' + k[1] + '需求编号' + str(k[2]) + '\n', '需求详细信息：', req_detail.to_json(force_ascii=False, orient='split', index=False), '相似需求如下：')
            # print('result_dict[k]111' * 5, result_dict[k])
            result_dict[k] = sorted(result_dict[k].items(), key=lambda x: x[1], reverse=True)
            # print('result_dict[k]222' * 5, result_dict[k])
            for ki in result_dict[k]:
                sim_detail = raw_info[(raw_info.batch == ki[0][0]) & (raw_info.projectNo == ki[0][1]) & (raw_info.requirementNO == ki[0][2])]
                print('相似需求批次' + ki[0][0] + '项目编号' + ki[0][1] + '需求编号' + str(ki[0][2]) + '\n', sim_detail.to_json(force_ascii=False, orient='split', index=False))
                print('相似度为：', ki[1])
    else:
        print('未发现相似' + flag + '需求')


def result_persist(result_dict, raw_info, new, related_no, flag):
    # print('right in result_persist  ' * 20)

    def df2dict(req_df):
        req_ser = req_df.iloc[0]
        req_dict = req_ser.to_dict()
        des = OrderedDict([
            ('batch', '需求批次'),
            ('projectNo', '项目编号'),
            ('project_name', '项目名称'),
            ('requirementNO', '需求编号'),
            ('requirement_name', '需求名称'),
            ('days_spent', '实际消耗时间'),
        ])
        if flag == 'cosmic':
            des['advocator'] = '需求提出人'
            des['requirement_detail'] = '需求详细信息'
            des['coding_requirement'] = '代码开发信息'
        else:
            des['work_cat'] = '工作类别'
            des['work_name'] = '工作名称'
            des['work_detail'] = '工作详情'
        for n in des:
            c = des[n] + '：' + str(req_dict[n])
            # print('here ' * 50 + '\n',c , '\n length:', len(c))
            if len(c) <= 100:
                des[n] = c
            else:
                des[n] = c[:35] + '\t ...'
                # print('len(c) > 100 ' * 20, des[n])
        # return '\n'.join(des.values())
        return des

    def in_paragraph(req_dict, DOCobj, sim_value=''):
        if not sim_value:
            mes = '#' * 50 + '\n' + '原需求如下：\n'
        else:
            mes = '相似需求如下：\n'
        p = DOCobj.add_paragraph()
        p.add_run(mes).bold = True
        mes = '\n'.join([req_dict[key] for key in ['batch', 'projectNo', 'project_name']]) + '\n'
        p.add_run(mes)
        mes = '\n'.join([req_dict[key] for key in ['requirementNO', 'requirement_name']]) + '\n'
        pr2 = p.add_run(mes)
        pr2.font.color.rgb = RGBColor(0xff, 0, 0)
        [req_dict.pop(key) for key in ['batch', 'projectNo', 'project_name', 'requirementNO', 'requirement_name']]
        mes = '\n'.join(req_dict.values()) + '\n'
        p.add_run(mes)
        if sim_value:
            p.add_run('相似度为：' + sim_value).font.color.rgb = RGBColor(255, 127, 0)

    if result_dict:
        doc = Document()
        doc.styles['Normal'].font.name = '宋体'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        doc.add_heading('TF-IDF Similarity Analyse', level=0)
        doc.add_heading(flag + '的相似需求情况如下：')
        pt = doc.add_paragraph()
        pt.add_run('共计项目个数：' + str(project_no) + '\n'
                   '共计需求个数：' + str(req_no) + '\n'
                   '发现有相似关系的需求：' + str(related_no) + '\n')
        pt.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for k in result_dict.keys():
            # mes += '发现需求批次：{}\t项目编号：{}\t需求编号：{}'.format(k[0], k[1], str(k[2]))
            req_detail = new[(new.batch == k[0]) & (new.projectNo == k[1]) & (new.requirementNO == k[2])]
            req_detail_dict = df2dict(req_detail)
            in_paragraph(req_detail_dict, doc)

            # mes = mes + '#' * 100 + '\n' + json.dumps(req_detail_des, ensure_ascii=False)
            result_dict[k] = sorted(result_dict[k].items(), key=lambda x: x[1], reverse=True)
            sim_req_detail = []
            # print(result_dict[k])
            for t in result_dict[k]:
                sim_req = raw_info[(raw_info.batch == t[0][0]) & (raw_info.projectNo == t[0][1]) & (raw_info.requirementNO == t[0][2])]
                sim_req = df2dict(sim_req)
                in_paragraph(sim_req, doc, str(t[1]))

            # mes = mes + '\n'.join(sim_req_detail) + '\n'
        # print("in_result_persist_3" * 20 + flag + "\n", mes)
        # doc.add_paragraph(mes)
        doc.save('./Doc/' + flag + '_TF-IDF_Q3.docx')
    else:
        print('未发现相似需求' + flag)


if __name__ == '__main__':

    raw_cosmic = pd.read_pickle('./data/cosmic_info_Q3.pkl').fillna('') ##########.sample(n=10)
    raw_noncosmic = pd.read_pickle('./data/noncosmic_info_Q3.pkl').fillna('')

    SUCCESS, info = get_corpus_data(raw_cosmic, raw_noncosmic, 'total')

    if not SUCCESS:
        print(info)
    # info.to_csv('./data/info.csv', encoding="utf_8_sig", index=False)
    dictionary = corpora.Dictionary(info.combo.to_list())

    ####################Total number of corpus positions (number of processed words)
    print('dictionary.num_pos'*5, dictionary.num_pos)
    dictionary.filter_extremes(no_below=1, no_above=0.9, keep_n=4000)

    ####################################################check the dictionary###########################

    print('len(dictionary.cfs)' * 5, len(dictionary.cfs))

    # # sum of the number of unique words per documen over the entire corpus
    # print('dictionary.num_nnz' * 5, dictionary.num_nnz)

    cfs = dictionary.cfs
    token_fs = {}
    for t, id in dictionary.token2id.items():
        # print(t, id)
        token_fs[t] = cfs[id]
    token_fs = sorted(token_fs.items(), key=lambda x:x[1], reverse=True)
    print('token_fs' * 10 + '\n', token_fs)
    ####################################################################################################

    dictionary.save_as_text('./data/total.dic')

    cosmic_sim_analyse = []
    noncosmic_sim_analyse = []
    project_no = 0
    req_no = 0
    # noncosmic_project_no = 0
    # noncosmic_req_no = 0
    for FLAG in ['cosmic', 'noncosmic']:
        fraud_new_req = get_new_req(raw_cosmic, raw_noncosmic, flag=FLAG)
        if fraud_new_req.empty:
            print('ERROR：新读入的' + FLAG + '需求为空')
            continue
        else:

            if FLAG == 'cosmic':
                raw_info = raw_cosmic
                project_no = len(set(fraud_new_req['projectNo']))
                req_no = fraud_new_req.shape[0]
                SUCCESS, new_req_corpus = get_corpus_data(cosmic_info=fraud_new_req, result_type=FLAG)
            else:
                raw_info = raw_noncosmic
                project_no = len(set(fraud_new_req['projectNo']))
                req_no = fraud_new_req.shape[0]
                SUCCESS, new_req_corpus = get_corpus_data(noncosmic_info=fraud_new_req, result_type=FLAG)
            if SUCCESS:
                solo = info[FLAG].apply(lambda x:np.NAN if len(x) == 0 else x).dropna()
                solo = info.iloc[solo.index][['batch', 'projectNo', 'requirementNO',FLAG]]
                sim_result, req_related_no = get_sim(solo, new_req_corpus, flag=FLAG)
                result_persist(sim_result, raw_info, fraud_new_req, req_related_no, FLAG)
            else:
                continue
