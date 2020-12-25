#-*- coding: utf-8 -*-
# @Time    : 2020/8/4 10:28
# @Author  : Nivina
# @Email   : nivinanull@163.com
# @File    : similarity_analyse.py

import pandas as pd
import datetime
import sys
import jieba
from collections import OrderedDict
import jieba.posseg as pseg
import re
from gensim import corpora, models
from docx import Document
from docx.shared import Pt,RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from gensim.similarities import SparseMatrixSimilarity


pd.set_option('display.max_columns',20) #给最大列设置为20列
pd.set_option('display.max_rows',10)#设置最大可见10行

# ##################################################最新季度信息
LATEST_BATCH = '2020Q4'  ##### 筛选出新需求
IF_MTR = 0


def get_all_reqs():
    raw_cosmic = pd.read_pickle('./data/20201218/cosmic_info.pkl').fillna('')  ##########.sample(n=10)
    # raw_cosmic = raw_cosmic[
    #     ['project_name', 'requirement_name', 'requirement_detail', 'advocator', 'coding_requirement']].astype(str)
    raw_noncosmic = pd.read_pickle('./data/20201218/noncosmic_info.pkl').fillna('')
    # raw_noncosmic = raw_noncosmic[['project_name', 'requirement_name', 'work_name', 'work_detail']].astype(str)
    if IF_MTR:
        maintenance_req = pd.read_pickle('./data/maintenance_req.pkl').fillna('')
        # maintenance_req = maintenance_req[['project_name', 'requirement_name', 'requirement_detail', 'advocator', 'phase_detail']].astype(str)
    else:
        maintenance_req = pd.DataFrame(
            columns=['type', 'batch', 'projectNO', 'requirementNO', 'file_trail', 'project_name', 'requirement_name', 'requirement_detail', 'advocator', 'phase_detail'])
    return raw_cosmic, raw_noncosmic, maintenance_req


def get_corpus_data(cosmic_info=pd.DataFrame(), noncosmic_info=pd.DataFrame(), maintenance_info=pd.DataFrame()):
    # 过滤词属性
    stop_property = ['nr', 'o', 'p', 'q', 'r', 'x', 'y', 'w', 'uj', 'm', 'un', 'c']
    stop_words_check = dict.fromkeys(stop_property)

    # 过滤长度为一的词
    stop_words_check['one'] = ['one']

    def num_process(w):
        if len(w.word) == 1:
            stop_words_check['one'].append(w.word)
            return False
        if w.flag in stop_property:
            if not stop_words_check[w.flag]:
                stop_words_check[w.flag] = [w.flag]
            else:
                stop_words_check[w.flag].append(w.word)
            return False
        else:
            return True

    def segment(row):
        if row:
            year_list = [str(i + datetime.datetime.now().year) for i in range(-20, 21)]
            year_related = ['\d*' + y + '\d*' for y in year_list]
            year_related_str = '|'.join(year_related)
            row = re.sub(year_related_str + '|\d\.|\d、|\d,|\d，|\W+', ' ', row).replace('_', ' ')
            jieba.load_userdict('./data/my_dict.txt')
            words = pseg.lcut(row)
            words = [i.word for i in words if num_process(i) and not re.search('R|r[0-9]\d+', i.word)]
        else:
            words = []
        return words

    def to_str(attri_row):
        joint_info = ''
        for c in range(len(attri_row.work_detail)):
            if not (attri_row.work_detail[c].strip() == '无' or attri_row.work_detail[c].strip() == ''):
                joint_info = attri_row.work_detail[c]
                if not (attri_row.work_name[c].strip() == '无' or attri_row.work_name[c].strip() == ''):
                    joint_info = joint_info + ' ' + attri_row.work_name[c]
        return joint_info

    def normalization(df, col_list):
        normal_cols = ['type', 'batch', 'projectNO', 'requirementNO', 'file_trail', 'joint_info']
        df['joint_info'] = ''
        for c in col_list:
            df['joint_info'] = df['joint_info'] + ' ' + df[c].astype(str)
        df['joint_info'] = df['joint_info'].map(segment)
        df = df[normal_cols]

    cosmic_inner = cosmic_info.copy()
    noncosmic_inner = noncosmic_info.copy()
    maintenance_inner = maintenance_info.copy()

    # ###########################将不空的需求信息，选择相应的列信息构建语料，存储为joint_info列
    if not noncosmic_inner.empty:
        noncosmic_inner['work_name_detail'] = noncosmic_inner[['work_name', 'work_detail']].apply(to_str, axis=1)
        normalization(noncosmic_inner, ['work_name_detail', 'requirement_name', 'project_name'])
    else:
        return False, 'ERROR: noncosmic信息不能为空，请检查'

    if not cosmic_inner.empty:
        normalization(cosmic_inner, ['project_name', 'requirement_name', 'requirement_detail', 'coding_requirement'])
    else:
        return False, 'ERROR: cosmic信息不能为空，请检查'

    if not maintenance_info.empty:
        normalization(maintenance_inner, ['project_name', 'requirement_name', 'requirement_detail', 'phase_detail'])

        ################################################save the stopwords ##########################################
        for k in stop_words_check.keys():
            with open('./data/stopwords_folder/' + k + '.txt', 'wb') as f:
                f.write(str(stop_words_check[k]).encode('utf-8'))
        ##############################################################################################################
        return True, pd.concat([cosmic_inner, noncosmic_inner, maintenance_inner])
    elif not IF_MTR:
        return True, pd.concat([cosmic_inner, noncosmic_inner])
    else:
        return False, 'ERROR: maintenance_info信息不能为空，请检查'


# result = get_sim(info[FLAG], new_req_corpus, flag)
# 已存需求的分词词料，新需求的分词词料，FLAG
def get_sim(all_reqs, new_req):
    # 利用cosmic和非cosmic信息构建的Dictionary
    dictionary = corpora.Dictionary.load_from_text('./data/total.dic')
    corpus = [dictionary.doc2bow(text) for text in all_reqs['joint_info']]

    # 得到tfidf模型
    tfidf = models.TfidfModel(corpus)

    # 得到稀疏矩阵相似度计算模型
    index = SparseMatrixSimilarity(tfidf[corpus], 4000)

    # 储存相似度结果的双层嵌套词典， 格式为：{new_req_id1:{sim_req_id1:sim_value, ...}, ...}
    # 其中主键id为需求的type, batch, projectNO, requirementNO, file_trail信息组成的tuple
    sim_dict = {}

    # 先根据dictionary将new_req（新需求）信息构建words bow，再使用tfidf（TF-IDF模型）将其转化为TF-IDF向量
    # 将新需求的TF-IDF向量输入到 index（SparseMatrixSimilarity模型）得到相似度矩阵sim， 大小为 [len(new_req) * len(all_req)]
    sim = index[tfidf[[dictionary.doc2bow(t) for t in new_req['joint_info']]]]
    print('相似矩阵的大小为：', sim.shape)
    relation_exist = []
    for i in range(sim.shape[0]):
        key = (new_req['type'].iloc[i], new_req['batch'].iloc[i], new_req['projectNO'].iloc[i],
               new_req['requirementNO'].iloc[i], new_req['file_trail'].iloc[i])
        # print(key)
        value = {}
        for j in range(sim.shape[1]):
            if sim[i][j] >= 0.86:
                inner_key = (all_reqs['type'].iloc[j], all_reqs['batch'].iloc[j], all_reqs['projectNO'].iloc[j],
                             all_reqs['requirementNO'].iloc[j], all_reqs['file_trail'].iloc[j])
                if inner_key == key:
                    continue
                else:
                    ## current_relation是 关系对set
                    current_relation = {key, inner_key}
                    if current_relation not in relation_exist:
                        value[inner_key] = sim[i][j]
                        relation_exist.append(current_relation)
        if value:
            sim_dict[key] = value
    return sim_dict, len(relation_exist)


def result_persist(result_dict, reqs, related_no):

    def df2dict(req_df):
        req_ser = req_df.iloc[0]
        req_dict = req_ser.to_dict()
        des = OrderedDict([
            ('type', '需求类型'),
            ('batch', '需求批次'),
            ('projectNO', '项目编号'),
            ('project_name', '项目名称'),
            ('requirementNO', '需求编号'),
            ('requirement_name', '需求名称'),
            ('file_trail', '路径名称'),
            ('days_spent', '实际消耗时间'),
            ('advocator', '需求提出人'),
            ('requirement_detail', '需求详细信息'),
            ('coding_requirement', '代码开发信息'),
            ('work_cat', '工作类别'),
            ('work_name', '工作名称'),
            ('work_detail', '工作详情'),
            ('phase_detail', '任务完成说明'),
        ])
        for n in des:
            if req_dict[n] != '':
                content = des[n] + '：' + str(req_dict[n])
                if len(content) <= 100:
                    des[n] = content
                else:
                    des[n] = content[:35] + '\t ...'
        return des

    def in_paragraph(req_dict, DOCobj, sim_value=''):
        if not sim_value:
            mes = '#' * 50 + '\n' + '原需求如下：\n'
        else:
            mes = '相似需求如下：\n'
        p = DOCobj.add_paragraph()
        p.add_run(mes).bold = True
        mes = '\n'.join([req_dict[key] for key in ['type', 'batch', 'projectNO', 'project_name']]) + '\n'
        p.add_run(mes)
        mes = '\n'.join([req_dict[key] for key in ['requirementNO', 'requirement_name']]) + '\n'
        pr2 = p.add_run(mes)
        pr2.font.color.rgb = RGBColor(0xff, 0, 0)
        [req_dict.pop(key) for key in ['type', 'batch', 'projectNO', 'project_name', 'requirementNO', 'requirement_name']]
        mes = '\n'.join([content for content in req_dict.values() if content not in
                         ['路径名称', '需求提出人', '需求详细信息', '代码开发信息', '工作类别', '工作名称',  '工作详情', '任务完成说明']]) + '\n'
        p.add_run(mes)
        if sim_value:
            p.add_run('相似度为：' + sim_value).font.color.rgb = RGBColor(255, 127, 0)

    if result_dict:
        project_no = reqs['project_name'].nunique()
        req_no = reqs['requirementNO'].nunique()
        doc = Document()
        doc.styles['Normal'].font.name = '宋体'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        doc.add_heading('TF-IDF Similarity Analyse', level=0)
        # doc.add_heading(flag + '的相似需求情况如下：')
        pt = doc.add_paragraph()
        pt.add_run('共计项目个数：' + str(project_no) + '\n'
                   '共计需求个数：' + str(req_no) + '\n'
                   '发现有相似关系的需求：' + str(related_no) + '\n')
        pt.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for k in result_dict.keys():
            # mes += '发现需求批次：{}\t项目编号：{}\t需求编号：{}'.format(k[0], k[1], str(k[2]))
            req_detail = reqs[
                (reqs.type == k[0]) & (reqs.batch == k[1]) & (reqs.projectNO == k[2]) & (reqs.requirementNO == k[3]) & (reqs.file_trail == k[4])]
            req_detail_dict = df2dict(req_detail)
            in_paragraph(req_detail_dict, doc)

            result_dict[k] = sorted(result_dict[k].items(), key=lambda x: x[1], reverse=True)
            sim_req_detail = []
            print('key ' * 8 + '\n', k, '\n' + 'result_dict[k] ' * 8 + '\n', result_dict[k])
            for t in result_dict[k]:
                sim_req = reqs[(reqs.type == t[0][0]) & (reqs.batch == t[0][1]) & (reqs.projectNO == t[0][2]) & (
                            reqs.requirementNO == t[0][3]) & (reqs.file_trail == t[0][4])]
                sim_req = df2dict(sim_req)
                in_paragraph(sim_req, doc, str(t[1]))

        doc.save('./Doc/TF-IDF_' + LATEST_BATCH + '.docx')
    else:
        print('未发现相似需求')


if __name__ == '__main__':

    cosmic_reqs, noncosmic_reqs, maintenance_reqs = get_all_reqs()
    all_for_query = pd.concat([cosmic_reqs, noncosmic_reqs, maintenance_reqs])

    IF_SUCCESS, result = get_corpus_data(cosmic_reqs, noncosmic_reqs, maintenance_reqs)

    if not IF_SUCCESS:
        print(result)
        sys.exit()
    all_reqs_in_words = result
    dictionary = corpora.Dictionary(all_reqs_in_words['joint_info'].to_list())

    # ####################Total number of corpus positions (number of processed words)即处理的文字总数
    print('dictionary.num_pos ' * 5 + '\n', dictionary.num_pos)
    dictionary.filter_extremes(no_below=2, no_above=0.8, keep_n=4000)

    # ####################################################check the dictionary###########################
    print('dictionary 词典大小 ' * 5, len(dictionary.cfs))
    cfs = dictionary.cfs
    token_fs = {}
    for t, id in dictionary.token2id.items():
        # print(t, id)
        token_fs[t] = cfs[id]
    token_fs = sorted(token_fs.items(), key=lambda x: x[1], reverse=True)
    print('token_fs' * 10 + '\n', token_fs)
    ####################################################################################################

    dictionary.save_as_text('./data/total.dic')

    cosmic_sim_analyse = []
    noncosmic_sim_analyse = []
    latest_reqs_in_words = all_reqs_in_words[all_reqs_in_words['batch'] == LATEST_BATCH]
    print('check the length of requirements ' * 2 + '\n', all_reqs_in_words.shape, latest_reqs_in_words.shape)
    sim_result, req_related_no = get_sim(all_reqs_in_words, latest_reqs_in_words)
    # print(sim_result)
    all_for_query = all_for_query.fillna('')
    result_persist(sim_result, all_for_query,  req_related_no)
